from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_mak_template():
    doc = Document()

    # --- 1. SETUP NARROW MARGINS ---
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # --- 2. SET DEFAULT FONT (Nirmala UI) ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Nirmala UI'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:cs'), 'Nirmala UI')

    # --- 3. HEADER (Table Layout) ---
    p_school = doc.add_paragraph("எம். ஏ.கே. குழுமப் பள்ளிகள்")
    p_school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_school.runs[0].bold = True
    p_school.runs[0].font.size = Pt(14)

    table = doc.add_table(rows=3, cols=2)
    table.autofit = True
    table.width = Inches(7.5)

    # Row 1: Class (Left) | Marks (Right)
    cell_1a = table.cell(0, 0)
    p = cell_1a.paragraphs[0]
    p.add_run("வகுப்பு: ").bold = True
    p.add_run("{{ header.class }}").bold = True
    
    cell_1b = table.cell(0, 1)
    p = cell_1b.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("மதிப்பெண்கள்: {{ header.marks }}").bold = True

    # Row 2: Date (Left) | Time (Right)
    cell_2a = table.cell(1, 0)
    p = cell_2a.paragraphs[0]
    p.add_run("தேதி: {{ header.date }}").bold = True

    cell_2b = table.cell(1, 1)
    p = cell_2b.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("நேரம்: {{ header.time }}").bold = True
    
    # Row 3: Subject
    cell_3a = table.cell(2, 0)
    p = cell_3a.paragraphs[0]
    p.add_run("தமிழ் முதல் தாள்").bold = True

    doc.add_paragraph("__________________________________________________________________________________________")

    # --- 4. MAK CONTENT LOOP ---
    doc.add_paragraph("{%p for section in sections %}")
    
    # Section Header
    p_sec = doc.add_paragraph()
    p_sec.add_run("{{ section.roman }}. {{ section.title }}")
    p_sec.runs[0].bold = True
    p_sec.paragraph_format.tab_stops.add_tab_stop(Inches(7.3), WD_ALIGN_PARAGRAPH.RIGHT)
    p_sec.add_run("\t{{ section.marks_eq }}").bold = True
    
    # Questions Loop
    doc.add_paragraph("{%p for q in section.questions %}")
    doc.add_paragraph("{{ q.no }}. {{ q.text }}")
    
    # MCQ Options
    doc.add_paragraph("{%p if q.options %}")
    p_opt = doc.add_paragraph()
    p_opt.paragraph_format.left_indent = Inches(0.4) 
    
    # Using specific Tamil labels (அ, ஆ, இ, ஈ)
    p_opt.add_run("{% if q.options[0] %} அ) {{ q.options[0] }}    {% endif %}")
    p_opt.add_run("{% if q.options[1] %} ஆ) {{ q.options[1] }}    {% endif %}")
    p_opt.add_run("{% if q.options[2] %} இ) {{ q.options[2] }}    {% endif %}")
    p_opt.add_run("{% if q.options[3] %} ஈ) {{ q.options[3] }}    {% endif %}")
    
    doc.add_paragraph("{%p endif %}")
    doc.add_paragraph("{%p endfor %}") # End Question
    doc.add_paragraph("{%p endfor %}") # End Section
    doc.add_paragraph("_______________________________________________")

    doc.save("template_mak.docx")
    print("✅ SUCCESS: 'template_mak.docx' created!")

def create_generic_template():
    doc = Document()
    
    # --- GENERIC CONTENT LOOP (For Choose/Para/Both) ---
    doc.add_paragraph("{%p for item in items %}")

    # Paragraph Logic
    doc.add_paragraph("{%p if item.type == 'para' %}")
    doc.add_paragraph("--------------------------------------------------")
    p = doc.add_paragraph("Topic: {{ item.heading }}")
    p.runs[0].bold = True
    doc.add_paragraph("{{ item.text }}")
    doc.add_paragraph("--------------------------------------------------")
    doc.add_paragraph("{%p endif %}")

    # MCQ Logic
    doc.add_paragraph("{%p if item.type == 'mcq' %}")
    doc.add_paragraph("{{ item.q_no }}. {{ item.text }}")
    doc.add_paragraph("    a) {{ item.options[0] }}")
    doc.add_paragraph("    b) {{ item.options[1] }}")
    doc.add_paragraph("    c) {{ item.options[2] }}")
    doc.add_paragraph("    d) {{ item.options[3] }}")
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph("{%p endfor %}")

    doc.save("template.docx")
    print("✅ SUCCESS: 'template.docx' created!")

if __name__ == "__main__":
    create_mak_template()
    create_generic_template()